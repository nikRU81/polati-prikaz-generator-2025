from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import os

app = Flask(__name__)

# Путь к логотипу
LOGO_PATH = os.path.join(os.path.dirname(__file__), 'static', 'img', 'logo.png')

@app.route('/')
def index():
    """Главная страница с формой"""
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_prikaz():
    """Генерация приказа на основе данных формы"""
    try:
        data = request.get_json()
        
        # Валидация данных
        if not all(k in data for k in ['day', 'month', 'year', 'orderNumber', 'orderTitle', 'preamble', 'punkts']):
            return jsonify({'error': 'Не все обязательные поля заполнены'}), 400
        
        if not data['punkts'] or len(data['punkts']) == 0:
            return jsonify({'error': 'Необходимо добавить хотя бы один пункт приказа'}), 400
        
        # Генерируем документ
        doc_buffer = create_prikaz_document(data)
        
        # Формируем имя файла
        filename = f"Приказ_ПОЛАТИ_{data['orderNumber'].replace('/', '-')}.docx"
        
        return send_file(
            doc_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        print(f"Ошибка при генерации: {str(e)}")
        return jsonify({'error': f'Ошибка при генерации документа: {str(e)}'}), 500

def create_prikaz_document(data):
    """Создание документа приказа согласно стандарту ПОЛАТИ 2025"""
    
    FONT_NAME = 'Times New Roman'
    FONT_SIZE_MAIN = 12
    
    # Создаем документ
    doc = Document()
    
    # Настройка полей страницы
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    
    # === HEADER С ЛОГОТИПОМ (только 1-я страница) ===
    section.different_first_page_header_footer = True
    first_page_header = section.first_page_header
    
    # Логотип
    if os.path.exists(LOGO_PATH):
        para_logo = first_page_header.paragraphs[0] if first_page_header.paragraphs else first_page_header.add_paragraph()
        para_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = para_logo.add_run()
        run_logo.add_picture(LOGO_PATH, height=Cm(1.2))
    
    first_page_header.add_paragraph()  # Пустая строка
    
    # РЕКВИЗИТЫ В ТАБЛИЦЕ 3x3
    table = first_page_header.add_table(rows=3, cols=3)
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Установка минимальной высоты строк
    for row in table.rows:
        tr = row._element
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '0')
        trHeight.set(qn('w:hRule'), 'auto')
        trPr.append(trHeight)
    
    # Строка 1
    cell = table.cell(0, 0)
    cell.width = Cm(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('ООО «ПОЛАТИ»')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    cell = table.cell(0, 1)
    cell.width = Cm(5.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('Тел: 8 (800) 234-22-77')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    cell = table.cell(0, 2)
    cell.width = Cm(5.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('ОГРН 1145029009982')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    # Строка 2
    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('141006, г. Мытищи, Московская')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    cell = table.cell(1, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('info@polati.ru')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    cell = table.cell(1, 2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('ИНН 5029188770')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    # Строка 3
    cell = table.cell(2, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('область, Олимпийский пр-т., стр. 29а,')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    cell = table.cell(2, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('polati.ru')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    cell = table.cell(2, 2)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('КПП 502901001')
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    
    # === ЗАГОЛОВОК "ПРИКАЗ" ===
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run('ПРИКАЗ')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    
    # === ТАБЛИЦА С ДАТОЙ И НОМЕРОМ (3 ячейки) ===
    table_date = doc.add_table(rows=1, cols=3)
    tbl = table_date._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Ячейка 1: День и месяц (слева)
    cell = table_date.rows[0].cells[0]
    cell.width = Cm(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run('«')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    run = p.add_run(data['day'])
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.underline = True
    
    run = p.add_run('» ')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    run = p.add_run(data['month'])
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.underline = True
    
    # Ячейка 2: Год (центр)
    cell = table_date.rows[0].cells[1]
    cell.width = Cm(5.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run(data['year'])
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.underline = True
    
    run = p.add_run(' г.')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    # Ячейка 3: Номер (справа)
    cell = table_date.rows[0].cells[2]
    cell.width = Cm(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    run = p.add_run('№ ')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    run = p.add_run(data['orderNumber'])
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.underline = True
    
    # === г. Мытищи ===
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run('г. Мытищи')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    # === НАЗВАНИЕ ПРИКАЗА ===
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(data['orderTitle'])
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    
    # === ПРЕАМБУЛА ===
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(data['preamble'])
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    # === ПРИКАЗЫВАЮ ===
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run('ПРИКАЗЫВАЮ:')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    
    # === ПУНКТЫ ПРИКАЗА ===
    for punkt in data['punkts']:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(f"{punkt['number']}. {punkt['text']}")
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
    
    # === ФИНАЛЬНЫЕ ПУНКТЫ ===
    last_punkt_num = len(data['punkts']) + 1
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f'{last_punkt_num}. Контроль исполнения настоящего приказа оставляю за собой.')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run(f'{last_punkt_num + 1}. Приказ вступает в силу с момента его подписания.')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    # === 5 ПУСТЫХ СТРОК ===
    for _ in range(5):
        doc.add_paragraph()
    
    # === ПОДПИСЬ ГД ===
    table = doc.add_table(rows=1, cols=3)
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    cell = table.rows[0].cells[0]
    cell.width = Cm(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Генеральный директор')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    cell = table.rows[0].cells[1]
    cell.width = Cm(5.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('__________________')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    cell = table.rows[0].cells[2]
    cell.width = Cm(6.0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('А.\u00A0А.\u00A0Зазыгин')
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    
    # === БЛОК ОЗНАКОМЛЕНИЯ ===
    if 'fios' in data and data['fios'] and len(data['fios']) > 0:
        doc.add_paragraph()
        doc.add_paragraph()
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run('С приказом ознакомлен(-ы):')
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = False
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Для каждого ФИО создаем строку
        for fio in data['fios']:
            table = doc.add_table(rows=1, cols=2)
            tbl = table._element
            tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Ячейка 1: ФИО
            cell = table.rows[0].cells[0]
            cell.width = Cm(2.5)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(fio)
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            
            # Ячейка 2: Линия и дата
            cell = table.rows[0].cells[1]
            cell.width = Cm(14.5)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.add_run('\u00A0' * 30)
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            
            run = p.add_run('_________________________________')
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            
            run = p.add_run('\u00A0')
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
            
            run = p.add_run('«__»_______20__г.')
            run.font.name = FONT_NAME
            run.font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Строка "Подпись"
        table = doc.add_table(rows=1, cols=2)
        tbl = table._element
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.add_tblPr()
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        cell = table.rows[0].cells[0]
        cell.width = Cm(2.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run('Подпись')
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        
        cell = table.rows[0].cells[1]
        cell.width = Cm(14.5)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        run = p.add_run('\u00A0' * 30)
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        
        run = p.add_run('_________________________________________')
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
    
    # Сохраняем в буфер
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
